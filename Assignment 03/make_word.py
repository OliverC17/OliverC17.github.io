from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

REPO_URL = "https://github.com/OliverC17/OliverC17.github.io"  
YOUR_NAME = "Zifan (Oliver) Chen"                      
COURSE_NAME = "AD 688 – Web Analytics"
ASSIGNMENT_NAME = "Module 3 Assignment – Big Data Visualization on Scale"


FIG_SECTION2 = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section2_boxplot.png"

FIG_SECTION3 = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section3_Bubble.png"

FIG_S4_ASSOC = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section4_ASSOC.png"
FIG_S4_BACH  = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section4_BACH.png"
FIG_S4_MAST  = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section4_master.png"
FIG_S4_PHD   = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section4_PhD.png"

FIG_S5_REMOTE = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section5_Remote.png"
FIG_S5_HYBRID = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section5_Hybrid.png"
FIG_S5_ONSITE = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section5_Onsite.png"
FIG_S5_HIST   = "/home/ubuntu/OliverC17.github.io/_site/Assignment 03/figures/Section5_HIST.png"
OUTPUT_DOCX = "AD688_Assignment03_Submission.docx"


SECTION2_TEXT = (
    "Key Insights:\n"
    "Across the top 15 NAICS2 industries, full-time positions (32+ hours) consistently demonstrate higher median starting "
    "salaries compared to both part-time categories, indicating that compensation premiums are strongly associated with "
    "full-time employment status.\n"
    "Industries such as Information, Finance, and Professional Services exhibit both higher central salary distributions and "
    "more extreme upper-end outliers, while sectors like Retail Trade and Other Services show comparatively lower medians and "
    "narrower interquartile ranges.\n"
    "Note: The box plot displays SALARY_FROM values (after removing missing and zero entries) grouped by NAICS2 industry and "
    "segmented by employment type."
)

SECTION3_TEXT = (
    "Key Insights:\n"
    "Among the top specialized occupations within the Business Intelligence domain, enterprise-level and platform-specific roles "
    "(e.g., ERP/Oracle/SAP) tend to exhibit higher median salaries, indicating compensation premiums for advanced or specialized "
    "expertise.\n"
    "Data Analyst roles show the highest posting volume but moderate median compensation, suggesting strong market demand with a "
    "broader entry-level representation.\n"
    "Note: Bubble size represents total job postings; the y-axis reflects the median SALARY_FROM based on non-null salary records."
)

S4_ASSOC_TEXT = (
    "Key Insights:\n"
    "For roles requiring an associate degree or lower, salaries are generally concentrated in the $40k–$90k range across most occupations.\n"
    "While salary tends to increase modestly with additional experience, the growth slope appears relatively moderate compared to higher education levels.\n"
    "Technical specialization (e.g., ERP/enterprise roles) can still command higher upper-bound salaries even at lower formal education levels."
)

S4_BACH_TEXT = (
    "Key Insights:\n"
    "Bachelor-level positions show a clearer positive relationship between experience and salary, with compensation increasingly clustering above $80k beyond 3–5 years.\n"
    "Specialized technical roles (e.g., Enterprise Architect, SAP/Oracle-related roles) demonstrate higher salary ceilings.\n"
    "Salary dispersion widens relative to the associate group, suggesting stronger differentiation by role complexity and skills."
)

S4_MAST_TEXT = (
    "Key Insights:\n"
    "Master’s degree roles show higher median salary levels and greater variability, with several occupations exceeding $150k at higher experience levels.\n"
    "Salary progression appears steeper in the early-to-mid career range (roughly 3–6 years), suggesting advanced education may accelerate compensation growth.\n"
    "High-end outliers indicate premium opportunities for specialized or strategic positions."
)

S4_PHD_TEXT = (
    "Key Insights:\n"
    "PhD-level roles exhibit the highest salary ceilings, though sample sizes tend to be smaller.\n"
    "Compensation is less tightly clustered and shows substantial dispersion, consistent with niche and senior-level roles.\n"
    "Experience remains important, while advanced credentials may enable access to higher-paying research-intensive or strategic positions."
)

S5_REMOTE_TEXT = (
    "Key Insights:\n"
    "Remote roles show salaries that generally increase with experience, with many mid-to-senior positions clustering between $90k and $150k.\n"
    "Platform-specific roles (e.g., Enterprise/ERP) appear more frequently at higher salary levels, indicating specialization premiums.\n"
    "Dispersion suggests remote opportunities span both mid-level and advanced analytical roles."
)

S5_HYBRID_TEXT = (
    "Key Insights:\n"
    "Hybrid roles appear limited in sample size compared to remote and onsite categories.\n"
    "Salary distribution is more concentrated, with fewer extreme high-end outliers, suggesting more standardized compensation bands for hybrid roles."
)

S5_ONSITE_TEXT = (
    "Key Insights:\n"
    "Onsite roles exhibit a broad salary distribution with visible upper-end outliers, indicating higher compensation potential for specialized or senior positions.\n"
    "Compared to remote roles, onsite positions show wider dispersion at higher experience levels.\n"
    "A dense cluster around $70k–$120k suggests this band represents a core compensation range."
)

S5_HIST_TEXT = (
    "Key Insights:\n"
    "Both remote and onsite roles share similar central salary ranges (roughly $60k–$120k).\n"
    "Onsite roles show a heavier right tail with more extreme high-end salaries, while remote roles appear more concentrated in mid-range salary bands.\n"
    "This suggests comparable baseline pay across arrangements, with onsite offering more extreme upper-bound outcomes."
)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)

def add_paragraph(doc: Document, text: str):
  
    for line in text.split("\n"):
        doc.add_paragraph(line)

def add_image(doc: Document, path: str, width_in: float = 6.5):
    if not os.path.exists(path):
        doc.add_paragraph(f"[Missing image file: {path}]")
        return
    doc.add_picture(path, width=Inches(width_in))


def main():
    doc = Document()

    # Title Page
    add_title(doc, ASSIGNMENT_NAME)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(f"{COURSE_NAME}\n").bold = True
    p.add_run(f"{YOUR_NAME}\n")
    p.add_run(datetime.now().strftime("%Y-%m-%d"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Submission / Repo URL
    add_heading(doc, "Repository", level=1)
    doc.add_paragraph("GitHub Repository HTTPS URL:")
    url_p = doc.add_paragraph(REPO_URL)
    url_p.runs[0].font.color.rgb = None  # keep default
    doc.add_paragraph("")

    # Section 2
    add_heading(doc, "Section 2 – Industry Salary Distribution (Box Plot)", level=1)
    add_image(doc, FIG_SECTION2)
    add_paragraph(doc, SECTION2_TEXT)

    doc.add_paragraph("")

    # Section 3
    add_heading(doc, "Section 3 – Specialized Occupation (Bubble Chart)", level=1)
    add_image(doc, FIG_SECTION3)
    add_paragraph(doc, SECTION3_TEXT)

    doc.add_paragraph("")

    # Section 4
    add_heading(doc, "Section 4 – Education Level (Scatter Plots)", level=1)

    add_heading(doc, "Associate or Lower", level=2)
    add_image(doc, FIG_S4_ASSOC)
    add_paragraph(doc, S4_ASSOC_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "Bachelor", level=2)
    add_image(doc, FIG_S4_BACH)
    add_paragraph(doc, S4_BACH_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "Master", level=2)
    add_image(doc, FIG_S4_MAST)
    add_paragraph(doc, S4_MAST_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "PhD", level=2)
    add_image(doc, FIG_S4_PHD)
    add_paragraph(doc, S4_PHD_TEXT)

    doc.add_page_break()

    # Section 5
    add_heading(doc, "Section 5 – Remote Work Type (Scatter + Histogram)", level=1)

    add_heading(doc, "Remote", level=2)
    add_image(doc, FIG_S5_REMOTE)
    add_paragraph(doc, S5_REMOTE_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "Hybrid", level=2)
    add_image(doc, FIG_S5_HYBRID)
    add_paragraph(doc, S5_HYBRID_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "Onsite", level=2)
    add_image(doc, FIG_S5_ONSITE)
    add_paragraph(doc, S5_ONSITE_TEXT)
    doc.add_paragraph("")

    add_heading(doc, "Salary Distribution Comparison (Histogram)", level=2)
    add_image(doc, FIG_S5_HIST)
    add_paragraph(doc, S5_HIST_TEXT)

    doc.save(OUTPUT_DOCX)
    print(f"✅ Saved: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()